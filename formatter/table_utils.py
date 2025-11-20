from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from .paragraph_utils import set_paragraph_format

def process_tables(doc):
    """Обработка всех таблиц в документе"""
    for table in doc.tables:
        # Высота строк
        for row in table.rows:
            row.height = Cm(0.8)
            for cell in row.cells:
                if cell.paragraphs:
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.left_indent = Cm(0)
                    p.paragraph_format.first_line_indent = Cm(0)
                    for run in p.runs:
                        run.font.size = Pt(12)
                        run.font.name = 'Times New Roman'
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Заголовок таблицы
        if table.rows:
            for cell in table.rows[0].cells:
                if cell.paragraphs:
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.bold = True

        # Отступы до/после таблицы
        tbl_element = table._element
        prev_sibling = tbl_element.getprevious()
        next_sibling = tbl_element.getnext()

        # Найти предыдущий параграф
        if prev_sibling is not None and prev_sibling.tag.endswith('p'):
            for para in doc.paragraphs:
                if para._element == prev_sibling:
                    # set_table_adjacent_format(para, after=False)
                    set_paragraph_format(
                        para,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        space_after=0,
                        space_before=6,
                        first_line_indent=Cm(0)
                    )
                    break

        # Найти следующий параграф
        if next_sibling is not None and next_sibling.tag.endswith('p'):
            for para in doc.paragraphs:
                if para._element == next_sibling:
                    # set_table_adjacent_format(para, after=True)
                    set_paragraph_format(
                        para,
                        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                        space_after=6,
                        space_before=0,
                        first_line_indent=Cm(0)
                    )
                    break