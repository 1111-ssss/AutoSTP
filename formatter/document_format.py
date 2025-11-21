# formatter/document_format.py
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .paragraph_utils import set_paragraph_format, process_paragraphs, format_paragraph
from .table_utils import process_tables
from .image_utils import process_images

def apply_base_style(doc):
    """Применяет базовый стиль Normal ко всему документу"""
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style_normal.paragraph_format
    pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_header(doc, fio, group, lab_number):
    """Добавляет колонтитул"""
    header = doc.sections[0].header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_text = f"{fio} {group} Лабораторная работа № {lab_number}"
    header_para.text = header_text
    set_paragraph_format(header_para, font_size=14, first_line_indent=Cm(0), alignment=WD_ALIGN_PARAGRAPH.CENTER)

def format_document_stp(input_path, output_path, fio, group, lab_number):
    doc = Document(input_path)
    apply_base_style(doc)
    add_header(doc, fio, group, lab_number)
    process_paragraphs(doc)
    process_tables(doc)
    process_images(doc)
    doc.save(output_path)
    print(f"Документ отформатирован и сохранён как: {output_path}")